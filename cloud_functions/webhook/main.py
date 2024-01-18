import functions_framework
import os
import re
import uuid
import json
import textract
import en_core_web_sm
import requests
from flask import Request, jsonify
from datetime import datetime
from google.cloud import bigquery
from googleapiclient.discovery import build
import google.auth
from google.oauth2.credentials import Credentials
import google.auth.transport.requests
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.http import MediaFileUpload
from google.cloud import aiplatform_v1
import vertexai
from vertexai.preview.language_models import TextEmbeddingModel
from vertexai.preview.language_models import TextGenerationModel
from docx import Document
from typing import List, Union, Dict, Optional, Tuple

def get_weighted_embeddings(chunk_embeddings: List[List[float]], chunk_lens: List[float]) -> List[float]:
    """
    Calculate the weighted average of embeddings based on chunk lengths.

    Args:
        chunk_embeddings (List[List[float]]): List of chunk embeddings, where each inner list represents
            the embeddings for a chunk.
        chunk_lens (List[float]): List of chunk lengths used as weights for the weighted average.

    Returns:
        List[float]: Weighted average embeddings based on chunk lengths.
    """
    # Calculate the sum of weights (chunk lengths)
    weights_sum = sum(chunk_lens)
    print("In weighted embeddings")
    # Calculate the weighted average for each element in the embeddings
    result = [
        sum(arr[i] * chunk_lens[k] / weights_sum for k, arr in enumerate(chunk_embeddings))
        for i in range(len(chunk_embeddings[0]))
    ]

    return result

def split_input(input_string: str, max_chunk_size: int) -> List[Dict[str, Union[str, int]]]:
    """
    Split the input string into chunks based on the specified maximum chunk size.

    Args:
        input_string (str): The input string to be split into chunks.
        max_chunk_size (int): The maximum size of each chunk.

    Returns:
        List[Dict[str, Union[str, int]]]: List of chunks, where each chunk is represented as a dictionary
            with keys 'chunk_content' (str) and 'chunk_size' (int).
    """
    chunks = []

    if len(input_string) > max_chunk_size:
        while len(input_string) > 0:
            # Find the last period (.) within the maximum chunk size
            last_period_index = input_string.rfind('.', 0, max_chunk_size)

            if last_period_index <= 0 or last_period_index > max_chunk_size:
                # If no period is found within the chunk size, split at the chunk size
                last_period_index = max_chunk_size

            # Extract the chunk
            chunk = input_string[:last_period_index + 1]

            # Remove the extracted chunk from the input
            input_string = input_string[last_period_index + 1:]

            # Trim leading and trailing whitespace
            chunk = chunk.strip()

            # Push the chunk to the result array
            chunks.append({
                "chunk_content": chunk,
                "chunk_size": len(chunk)
            })

    else:
        chunks = [{
            "chunk_content": input_string,
            "chunk_size": len(input_string)
        }]
    print("chunks are: ", chunks)
    return chunks

def generate_cover_letter(resume_text: str, job_text: str) -> Optional[str]:
    """
    Generate a cover letter for a job opening based on a candidate's resume.

    Args:
        resume_text (str): The text of the candidate's resume.
        job_text (str): The text of the job opening.

    Returns:
        Optional[str]: The generated cover letter text or None if an error occurs.
    """
    try:
        # Initialize Vertex AI
        vertexai.init(project="ml-spez-ccai", location="us-central1")

        # Parameters for text generation model
        parameters = {
            "max_output_tokens": 8192,
            "temperature": 0.0,
            "top_p": 0.95,
            "top_k": 40
        }

        # Load the text generation model
        model = TextGenerationModel.from_pretrained("text-bison-32k")

        # Create the prompt for text generation
        prompt = f"""Given a candidate's resume with text:
        {resume_text}
        Given a job opening with text:
        {job_text}
        Write a cover letter for the job opening on behalf of the candidate.
        Use relevant information from the candidate's resume to generate the letter,
        do not add fictional information."""

        # Generate the cover letter using the text generation model
        response = model.predict(prompt, **parameters)

        return response.text

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error: ' + e.message)
        else:
            print('Error: ' + str(e))
        return None

def delete_folders(folder_id: str) -> Optional[bool]:
    """
    Delete a folder from Google Drive based on its folder ID.

    Args:
        folder_id (str): The ID of the folder to be deleted.

    Returns:
        Optional[bool]: True if the folder is deleted successfully, None if an error occurs.
    """
    try:
        # Get credentials
        credentials = get_credentials()

        # Build Google Drive API service
        service = build('drive', 'v3', credentials=credentials)

        # Delete the folder using the Drive API
        service.files().delete(fileId=folder_id).execute()

        return True
    except HttpError as error:
        print(f"An error occurred: {error}")
        return None

def upload_file(file_name: str, folder_id: str, local_path: str) -> Optional[bool]:
    """
    Upload a file to Google Drive within a specified folder.

    Args:
        file_name (str): The name of the file to be uploaded.
        folder_id (str): The ID of the folder in which to upload the file.
        local_path (str): The local path of the file to be uploaded.

    Returns:
        Optional[bool]: True if the file is uploaded successfully, None if an error occurs.
    """
    try:
        # Get credentials
        credentials = get_credentials()

        # Build Google Drive API service
        service = build('drive', 'v3', credentials=credentials)

        # Set metadata for the file
        file_metadata = {
            'name': file_name,
            'parents': [folder_id]
        }

        # Create a media file upload instance
        media = MediaFileUpload(local_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # Upload the file using the Drive API
        file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()

        return True
    except HttpError as error:
        print(f"An error occurred: {error}")
        return None

def save_cl(text: str, file_name: str, folder_id: str) -> Optional[bool]:
    """
    Save a cover letter text to a Word document, and upload it to Google Drive.

    Args:
        text (str): The cover letter text to be saved.
        file_name (str): The desired name of the Word document (without extension).
        folder_id (str): The ID of the folder in which to upload the document.

    Returns:
        Optional[bool]: True if the document is saved and uploaded successfully, None if an error occurs.
    """
    try:
        # Append ".docx" to the file name
        file_name = file_name + ".docx"

        # Get the current timestamp
        current_timestamp = datetime.now().strftime("%d-%m-%Y %H:%M:%S")

        # Create a new Word document
        doc = Document()

        # Configure the header and footer
        section = doc.sections[0]
        header = section.header
        footer = section.footer
        h_paragraph = header.paragraphs[0]
        f_paragraph = footer.paragraphs[0]
        h_paragraph.text = "Welcome to Intelia ML Specialization Demo"
        f_paragraph.text = f"Generated at {current_timestamp}."

        # Add the cover letter text to the document
        doc.add_paragraph(text)

        # Save the document to a temporary location
        doc_path = "/tmp/" + file_name
        doc.save(doc_path)

        # Upload the document to Google Drive
        upload_success = upload_file(file_name, folder_id, doc_path)

        return upload_success
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def save_job(job_details: Dict[str, str], file_name: str, folder_id: str) -> Optional[bool]:
    """
    Save job details to a Word document and upload it to Google Drive.

    Args:
        job_details (Dict[str, str]): A dictionary containing job details.
        file_name (str): The desired name of the Word document (without extension).
        folder_id (str): The ID of the folder in which to upload the document.

    Returns:
        Optional[bool]: True if the document is saved and uploaded successfully, None if an error occurs.
    """
    try:
        # Append ".docx" to the file name
        file_name = file_name + ".docx"

        # Get the current timestamp
        current_timestamp = datetime.now().strftime("%d-%m-%Y %H:%M:%S")

        # Create a new Word document
        doc = Document()

        # Configure the header and footer
        section = doc.sections[0]
        header = section.header
        footer = section.footer
        h_paragraph = header.paragraphs[0]
        f_paragraph = footer.paragraphs[0]
        h_paragraph.text = "Welcome to Intelia ML Specialization Demo"
        f_paragraph.text = f"Generated at {current_timestamp}."

        # Add job title as heading
        doc.add_heading(job_details["title"], 0)

        # Add job details section heading
        doc.add_heading('Job Details', level=4)

        # Add job details to the document
        type_para = doc.add_paragraph()
        type_para.add_run('Work Type: ').bold = True
        type_para.add_run(job_details["formatted_work_type"])

        location_para = doc.add_paragraph()
        location_para.add_run('Location: ').bold = True
        location_para.add_run(job_details["location"])

        if job_details["min_salary"]:
            min_salary_para = doc.add_paragraph()
            min_salary_para.add_run('Minimum Salary: ').bold = True
            min_salary_para.add_run(str(job_details["min_salary"]))

        if job_details["max_salary"]:
            max_salary_para = doc.add_paragraph()
            max_salary_para.add_run('Maximum Salary: ').bold = True
            max_salary_para.add_run(str(job_details["max_salary"]))

        if job_details["pay_period"]:
            pay_period_para = doc.add_paragraph()
            pay_period_para.add_run('Pay Period: ').bold = True
            pay_period_para.add_run(job_details["pay_period"])

        stats_para = doc.add_paragraph()
        stats_text = f"{str(job_details.get('views', 0))} Views {str(job_details.get('applies', 0))} Applies"
        stats_para.add_run(stats_text)

        # Add job description section heading
        doc.add_heading('Job Description', level=4)

        # Add job description to the document
        doc.add_paragraph(job_details["description"])

        # Save the document to a temporary location
        doc_path = "/tmp/" + file_name
        doc.save(doc_path)

        # Upload the document to Google Drive
        upload_success = upload_file(file_name, folder_id, doc_path)

        return upload_success
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def get_job(job_id: str) -> Optional[Dict[str, str]]:
    """
    Retrieve job details from BigQuery based on the provided job ID.

    Args:
        job_id (str): The ID of the job to retrieve.

    Returns:
        Optional[Dict[str, str]]: A dictionary containing job details or None if an error occurs.
    """
    try:
        # Retrieve BigQuery table ID from environment variable
        jobs_table_id = os.environ.get("JOBS_TABLE_ID")

        # Create a BigQuery client
        client = bigquery.Client()

        # Construct the SQL query
        query = f'''
        SELECT
        job_id,
        title,
        formatted_work_type,
        description,
        max_salary,
        min_salary,
        pay_period,
        views,
        applies,
        location,
        job_posting_url
        FROM
        `{jobs_table_id}`
        WHERE
        job_id = "{job_id}"
        '''

        # Execute the query
        query_job = client.query(query)
        result_data = []

        # Process the query results
        results = query_job.result()  # Waits for the job to complete.
        for result in results:
            result_data.append(dict(result))

        # Return the first result (if any)
        return result_data[0] if result_data else None

    except Exception as e:
        if hasattr(e, 'message'):
            print('Unable to get BigQuery results: ' + e.message)
        else:
            print('Unable to get BigQuery results: ' + str(e))
        return None


def get_job_details(matches: Dict[str, float]) -> Optional[List[Dict[str, str]]]:
    """
    Retrieve job details from BigQuery based on a dictionary of job matches.

    Args:
        matches (Dict[str, float]): A dictionary mapping job IDs to match percentages.

    Returns:
        Optional[List[Dict[str, str]]]: A list of dictionaries containing job details, sorted by match percentage,
                                         or None if an error occurs.
    """
    try:
        # Extract job IDs from the matches dictionary
        job_ids = list(matches.keys())

        # Retrieve BigQuery table ID from environment variable
        jobs_table_id = os.environ.get("JOBS_TABLE_ID")

        # Create a BigQuery client
        client = bigquery.Client()

        # Construct the SQL query using UNNEST to filter by job IDs
        query = f'''
        SELECT
        job_id,
        title,
        formatted_work_type,
        max_salary,
        min_salary,
        pay_period,
        location
        FROM
        `{jobs_table_id}`
        WHERE
        job_id IN UNNEST({job_ids})
        '''

        # Execute the query
        query_job = client.query(query)
        result_data = []

        # Process the query results
        results = query_job.result()  # Waits for the job to complete.
        for result in results:
            result_dict = dict(result)

            # Calculate and add match percentage to the result dictionary
            match_percent = round(float(matches[result_dict["job_id"]]) * 100, 2)
            result_dict["match_percent"] = match_percent

            result_data.append(result_dict)

        # Sort the results by match percentage in descending order
        sorted_results = sorted(result_data, key=lambda x: x['match_percent'], reverse=True)

        return sorted_results

    except Exception as e:
        if hasattr(e, 'message'):
            print('Unable to get BigQuery results: ' + e.message)
        else:
            print('Unable to get BigQuery results: ' + str(e))
        return None


def get_matches(vector: List[float]) -> Dict[str, float]:
    """
    Get job matches based on a vector using a matching engine.

    Args:
        vector (List[float]): The input vector for which matches are to be found.

    Returns:
        Dict[str, float]: A dictionary mapping job IDs to match distances.
    """
    try:
        print("in get matches")
        # Retrieve match threshold from environment variable
        match_threshold = float(os.environ.get("MATCH_THRESHOLD"))

        # Set variables for the current deployed index.
        API_ENDPOINT="1782564241.us-central1-917573008156.vdb.vertexai.goog"
        INDEX_ENDPOINT="projects/917573008156/locations/us-central1/indexEndpoints/8350381794633187328"
        DEPLOYED_INDEX_ID="job_posting_deployed_index"

        # Configure Vector Search client
        client_options = {
        "api_endpoint": API_ENDPOINT
        }
        vector_search_client = aiplatform_v1.MatchServiceClient(
        client_options=client_options,
        )

        # Build FindNeighborsRequest object
        datapoint = aiplatform_v1.IndexDatapoint(
        feature_vector=vector
        )
        query = aiplatform_v1.FindNeighborsRequest.Query(
        datapoint=datapoint,
        # The number of nearest neighbors to be retrieved
        neighbor_count=10
        )
        request = aiplatform_v1.FindNeighborsRequest(
        index_endpoint=INDEX_ENDPOINT,
        deployed_index_id=DEPLOYED_INDEX_ID,
        # Request can have multiple queries
        queries=[query],
        return_full_datapoint=False,
        )

        # Execute the request
        response = vector_search_client.find_neighbors(request)
        print("vector search response: ", response)
        # Extract matches from the response
        matches = {}

        if len(response[0]) > 0:
            for id, neighbor in enumerate(response[0]):
                if neighbor.distance >= match_threshold:
                    matches[neighbor.id] = neighbor.distance
        print("matches are: ", matches)
        return matches

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in finding matches: ' + e.message)
        else:
            print('Error in finding matches: ' + str(e))
        return {}

def get_text_embedding(text: str) -> List[float]:
    """
    Get text embedding using a Large Language Model.

    Args:
        text (str): The input text for which embedding is to be obtained.

    Returns:
        List[float]: A list representing the text embedding.
    """
    try:
        # Initialize the Text Embedding Model
        model = TextEmbeddingModel.from_pretrained("textembedding-gecko")

        # Get embeddings for the input text
        embeddings = model.get_embeddings([text])

        # Extract the vector from the embeddings
        for embedding in embeddings:
            vector = embedding.values

        return vector

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in getting text embedding: ' + e.message)
        else:
            print('Error in getting text embedding: ' + str(e))
        return []

def get_token_count(content: str, model: str) -> Optional[int]:
    """
    Get the token count for the given content using a specified language model.

    Args:
        content (str): The input text for which token count is to be obtained.
        model (str): The name of the language model (e.g., "textembedding-gecko").

    Returns:
        Optional[int]: The total token count if successful, None if an error occurs.
    """
    try:
        print("in token count")
        if model == "textembedding-gecko":
            body = {
                "instances": [
                    {"content": content}
                ],
            }
        else:
            body = {
                "instances": [
                    {"prompt": content}
                ],
            }

        request_body = json.dumps(body)

        # Construct the endpoint URL
        endpoint = f"https://us-central1-aiplatform.googleapis.com/v1/projects/ml-spez-ccai/locations/us-central1/publishers/google/models/{model}:countTokens"

        # Get the access token
        access_token = get_default_token()

        # Construct the Authorization header
        auth = "Bearer " + access_token

        # Create the headers with the Authorization header
        headers = {
            'Authorization': auth,
            'Content-Type': 'application/json; charset=utf-8'
        }

        # Send the POST request with JSON data
        response = requests.post(endpoint, headers=headers, data=request_body)

        if response.status_code == 200:
            response_json = response.json()
            print("token count: ", int(response_json["totalTokens"]))
            return int(response_json["totalTokens"])
        else:
            print(f'POST request failed with status code {response.status_code}')
            return None

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in getting token count: ' + e.message)
        else:
            print('Error in getting token count: ' + str(e))
        return None


def get_default_token() -> Optional[str]:
    """
    Get the default access token for Google Cloud Platform credentials.

    Returns:
        Optional[str]: The access token if successful, None if an error occurs.
    """
    try:
        # Define the required credential scopes
        CREDENTIAL_SCOPES = ["https://www.googleapis.com/auth/cloud-platform"]

        # Get default credentials and project ID
        credentials, project_id = google.auth.default(scopes=CREDENTIAL_SCOPES)

        # Create a request object
        request = google.auth.transport.requests.Request()

        # Refresh the credentials to obtain a new access token
        credentials.refresh(request)

        # Extract the access token
        access_token = credentials.token

        return access_token

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in getting default token: ' + e.message)
        else:
            print('Error in getting default token: ' + str(e))
        return None

def get_sentences(text: str) -> Optional[str]:
    """
    Tokenize the input text into sentences using a natural language processing library.

    Args:
        text (str): The input text to be tokenized into sentences.

    Returns:
        Optional[str]: The tokenized sentences joined into a single string, or None if an error occurs.
    """
    try:
        # Replace newline characters with spaces
        text = text.replace('\n', ' ').replace('\r', '')

        # Load the English language model for spaCy
        nlp = en_core_web_sm.load()

        # Process the text with spaCy's NLP pipeline
        doc = nlp(text)

        # Extract sentences from the processed document
        return_sentences = []

        for sent in doc.sents:
            string_sentence = str(sent)
            string_sentence = re.sub("\s\s+", " ", string_sentence)

            if string_sentence.strip() != "":
                return_sentences.append(string_sentence)

        # Join the sentences into a single string with newline separators
        return_text = '\n'.join(return_sentences)

        return return_text

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in getting sentences: ' + e.message)
        else:
            print('Error in getting sentences: ' + str(e))
        return None

def get_txt_pdf(path: str) -> Optional[str]:
    """
    Extract text content from a PDF file using textract.

    Args:
        path (str): The file path of the PDF document.

    Returns:
        Optional[str]: The extracted text content, or None if an error occurs.
    """
    try:
        # Use textract to extract text from the PDF file
        text = textract.process(path)
        text = text.decode("utf8")
        return text

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in extracting text from PDF: ' + e.message)
        else:
            print('Error in extracting text from PDF: ' + str(e))
        return None

def get_txt_docx(path: str) -> Optional[str]:
    """
    Extract text content from a DOCX file using the python-docx library.

    Args:
        path (str): The file path of the DOCX document.

    Returns:
        Optional[str]: The extracted text content, or None if an error occurs.
    """
    try:
        # Open the DOCX file using python-docx
        doc = Document(path)

        # Extract text from paragraphs
        text = [paragraph.text for paragraph in doc.paragraphs]

        # Join the text into a single string with newline separators
        return '\n'.join(text)

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in extracting text from DOCX: ' + e.message)
        else:
            print('Error in extracting text from DOCX: ' + str(e))
        return None

def download_file(folder_id: str, file_name: str) -> Optional[str]:
    """
    Download a file from Google Drive given the folder ID and file name.

    Args:
        folder_id (str): The ID of the Google Drive folder containing the file.
        file_name (str): The name of the file to be downloaded.

    Returns:
        Optional[str]: The local file path where the file is downloaded, or None if an error occurs.
    """
    try:
        # Get credentials for Google Drive API
        credentials = get_credentials()

        # Build the Google Drive API service
        service = build('drive', 'v3', credentials=credentials)

        # List files in the specified folder
        response = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields='files(id, name)'
        ).execute()

        # Find the file ID matching the specified file name
        files = response.get('files', [])
        file_id = None
        for file in files:
            if file["name"] == file_name:
                file_id = file["id"]

        if not file_id:
            print(f"File '{file_name}' not found in folder with ID '{folder_id}'.")
            return None

        # Download the file
        request = service.files().get_media(fileId=file_id)
        file_path = os.path.join("/tmp/", file_name)
        with open(file_path, 'wb') as file:
            downloader = MediaIoBaseDownload(file, request)
            done = False
            while done is False:
                _, done = downloader.next_chunk()

        return file_path

    except HttpError as error:
        print(f"An error occurred: {error}")
        return None

def get_folder_contents(folder_id: str) -> Optional[List[Dict[str, str]]]:
    """
    Get the contents of a Google Drive folder given its ID.

    Args:
        folder_id (str): The ID of the Google Drive folder.

    Returns:
        Optional[List[Dict[str, str]]]: A list of dictionaries containing file details (id, name),
        or None if an error occurs.
    """
    try:
        # Get credentials for Google Drive API
        credentials = get_credentials()

        # Build the Google Drive API service
        service = build('drive', 'v3', credentials=credentials)

        # List files in the specified folder
        response = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields='files(id, name)'
        ).execute()

        # Extract file details from the response
        files = response.get('files', [])

        return files

    except HttpError as error:
        print(f"An error occurred: {error}")
        return None


def create_folder(name: str, root: str) -> Optional[Tuple[str, str]]:
    """
    Create a new folder in Google Drive with the specified name and parent folder.

    Args:
        name (str): The name of the new folder.
        root (str): The ID of the parent folder where the new folder will be created.

    Returns:
        Optional[Tuple[str, str]]: A tuple containing the ID of the created folder and
        a public link to the folder, or None if an error occurs.
    """
    try:
        # Get credentials for Google Drive API
        credentials = get_credentials()

        # Build the Google Drive API service
        service = build('drive', 'v3', credentials=credentials)

        # Create metadata for the new folder
        file_metadata = {
            "name": name,
            "parents": [root],
            "mimeType": "application/vnd.google-apps.folder",
        }

        # Create the new folder
        file = service.files().create(body=file_metadata, fields="id").execute()
        folder_id = file.get("id")

        # Set permissions for public sharing
        permission = {
            'type': 'anyone',
            'role': 'writer'
        }
        permission_result = service.permissions().create(
            fileId=folder_id,
            body=permission,
            fields='id'
        ).execute()

        # Generate a public link to the folder
        public_link = f'https://drive.google.com/drive/folders/{folder_id}?usp=sharing'

        return folder_id, public_link

    except HttpError as error:
        print(f"An error occurred: {error}")
        return None

def get_credentials() -> Optional[Credentials]:
    """
    Retrieve Google API credentials for the specified scopes.

    Returns:
        Optional[Credentials]: Google API credentials, or None if an error occurs.
    """
    try:
        # Define the desired Google API scopes
        CREDENTIAL_SCOPES = ["https://www.googleapis.com/auth/cloud-platform", "https://www.googleapis.com/auth/drive"]

        # Use the google.auth.default function to get credentials
        credentials, project_id = google.auth.default(scopes=CREDENTIAL_SCOPES)

        return credentials

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in obtaining Google API credentials: ' + e.message)
        else:
            print('Error in obtaining Google API credentials: ' + str(e))
        return None

def watch_changes(folder_id: str) -> None:
    """
    Set up a watch to receive notifications about changes in a Google Drive folder.

    Args:
        folder_id (str): The ID of the Google Drive folder to watch.

    Returns:
        None
    """
    try:
        # Get credentials for Google Drive API
        credentials = get_credentials()

        # Build the Google Drive API service
        service = build('drive', 'v3', credentials=credentials)

        # Create a watch request for the folder
        watch_request = {
            'id': str(uuid.uuid4()),
            'type': 'web_hook',
            'address': 'https://us-central1-ml-spez-ccai.cloudfunctions.net/webhook',
            'payload': True
        }

        # Execute the watch request
        watch_response = service.files().watch(fileId=folder_id, body=watch_request).execute()

    except Exception as e:
        if hasattr(e, 'message'):
            print('Error in setting up watch request: ' + e.message)
        else:
            print('Error in setting up watch request: ' + str(e))

@functions_framework.http
def webhook(request: Request) -> jsonify:
	"""
	Handle incoming webhook requests.

	Args:
		request (Request): The incoming HTTP request.

	Returns:
		jsonify: JSON response based on the incoming request.
	"""
	# Get the root folder ID from environment variables
	root_folder_id = os.environ.get("ROOT_FOLDER_ID")
	if 'Content-Type' in request.headers and request.headers['Content-Type'] == 'application/json':
        # Parse JSON from the incoming request
		request_json = request.get_json(silent=True)
		if "sessionInfo" in request_json and "session" in request_json["sessionInfo"]:
			session_info = request_json["sessionInfo"]
			if "parameters" in session_info:
				session_parameters = session_info["parameters"]
			else:
				session_parameters = {}
			session_id_regex = r".+\/sessions\/(.+)"
			session = request_json["sessionInfo"]["session"]
			regex_match = re.search(session_id_regex, session)
			session_id = regex_match.group(1)
		if "pageInfo" in request_json and "formInfo" in request_json["pageInfo"] and "parameterInfo" in request_json["pageInfo"]["formInfo"]:
			page_parameters = request_json["pageInfo"]["formInfo"]["parameterInfo"]
		else:
			page_parameters = []
		if "fulfillmentInfo" in request_json and "tag" in request_json["fulfillmentInfo"] and session_id:
			tag = request_json["fulfillmentInfo"]["tag"]
			if tag == "init_folders":
				if "folders_created" not in session_parameters:
					session_folder_id, session_folder_link = create_folder(session_id, root_folder_id)
					resume_folder_id, resume_folder_link = create_folder("Resumes", session_folder_id)
					cl_folder_id, cl_folder_link = create_folder("Cover Letters", session_folder_id)
					matches_folder_id, matches_folder_link = create_folder("Matching Jobs", session_folder_id)
					json_response = {
						"sessionInfo": {
							"parameters": {
								"folders_created": True,
								"session_folder_id": session_folder_id,
								"session_folder_link": session_folder_link,
								"resume_folder_id": resume_folder_id,
								"resume_folder_link": resume_folder_link,
								"cl_folder_id": cl_folder_id,
								"cl_folder_link": cl_folder_link,
								"matches_folder_id": matches_folder_id,
								"matches_folder_link": matches_folder_link
							},
						},
					}
					return jsonify(json_response)
			if tag == "create_folder":
				if "folders_created" not in session_parameters:
					session_folder_id, session_folder_link = create_folder(session_id, root_folder_id)
					resume_folder_id, resume_folder_link = create_folder("Resumes", session_folder_id)
					cl_folder_id, cl_folder_link = create_folder("Cover Letters", session_folder_id)
					matches_folder_id, matches_folder_link = create_folder("Matching Jobs", session_folder_id)
					#watch_changes(folder_id)
					html =  f'''
					<p>Please use this folder to upload a copy of your resume.</p>
					<p><a href="{resume_folder_link}" target="_blank">Upload Your Resume</a></p>
					<p>Please click the button below once done.</p>
					'''
					json_response = {
						"sessionInfo": {
							"parameters": {
								"folders_created": True,
								"session_folder_id": session_folder_id,
								"session_folder_link": session_folder_link,
								"resume_folder_id": resume_folder_id,
								"resume_folder_link": resume_folder_link,
								"cl_folder_id": cl_folder_id,
								"cl_folder_link": cl_folder_link,
								"matches_folder_id": matches_folder_id,
								"matches_folder_link": matches_folder_link
							},
						},
						'fulfillment_response': {
							'messages': [
								{
									'payload': {
										'richContent': [
											[
												{
													"type": "html",
													"html": html
												},
												{
													"type": "chips",
													"options": [
													{
														"text": "I have uploaded the file"
													}
												]
												}
											]
										]
									}
								}
							]
						}
					}
				else:
					resume_folder_link = session_parameters["resume_folder_link"]
					html =  f'''
					<p>Please use this folder to upload a copy of your resume.</p>
					<p><a href="{resume_folder_link}" target="_blank">Upload Your Resume</a></p>
					<p>Please click the button below once done.</p>
					'''
					json_response = {
						'fulfillment_response': {
							'messages': [
								{
									'payload': {
										'richContent': [
											[
												{
													"type": "html",
													"html": html
												},
												{
													"type": "chips",
													"options": [
													{
														"text": "I have uploaded the file"
													}
												]
												}
											]
										]
									}
								}
							]
						}
					}
				return jsonify(json_response)
			if tag == "file_uploaded":
				resume_folder_id = session_parameters["resume_folder_id"]
				files = get_folder_contents(resume_folder_id)
				if len(files) == 0:
					text = "I was unable to find any files."
					resume_folder_link = session_parameters["resume_folder_link"]
					html =  f'''
					<p>Please use this folder to upload a copy of your resume.</p>
					<p><a href="{resume_folder_link}" target="_blank">Upload Your Resume</a></p>
					<p>Please click the button below once done.</p>
					'''
					json_response = {
							'fulfillment_response': {
								'messages': [
									{"text": {"text": [text]}},
									{
										'payload': {
											'richContent': [
												[
													{
														"type": "html",
														"html": html
													},
													{
														"type": "chips",
														"options": [
														{
															"text": "I have uploaded the file"
														}
													]
													}
												]
											]
										}
									}
								]
							}
						}
				else:
					options = []
					text = "Please click on the file name to process."
					for file in files:
						option_text = "Filename: "+file["name"]
						options.append({
							"type": "chips",
							"options": [
								{
								"text": option_text
								}
							]
						})
					json_response = {
							"page_info": {
										"form_info": {
											"parameter_info": [
												{
													"displayName": "files_displayed",
													"required": False,
													"state": "VALID",
													"value": True,
												},
											],
										},
									},
							'fulfillment_response': {
								'messages': [
									{"text": {"text": [text]}},
									{
										'payload': {
											'richContent': [options]
										}
									}
								]
							}
						}
				return jsonify(json_response)
			if tag == "file_confirmed":
				for parameter in page_parameters:
					if parameter["displayName"] == "files_displayed" and parameter["value"] == True:
						file_name = request_json["text"][10:]
						resume_folder_id = session_parameters["resume_folder_id"]
						file_path = download_file(resume_folder_id, file_name)
						if file_path and file_path.endswith(".docx"):
							text = get_txt_docx(file_path)
						if file_path and file_path.endswith(".pdf"):
							text = get_txt_pdf(file_path)
						if text:
							content = get_sentences(text)
							token_count = get_token_count(content,"textembedding-gecko")
							if token_count and token_count<3072:
								vector = get_text_embedding(text)
							else:
								print("input too big")
								chunks = split_input(content,10000)
								chunk_embeddings = []
								chunk_lengths = []
								for chunk in chunks:
									vector = get_text_embedding(chunk["chunk_content"])
									chunk_embeddings.append(vector)
									chunk_lengths.append(chunk["chunk_size"])
								vector = get_weighted_embeddings(chunk_embeddings,chunk_lengths)

							matches = get_matches(vector)
							job_details = get_job_details(matches)
							options = []
							text = ''
							for job in job_details:
								match_percent = job['match_percent']
								option_text = f"Export: {job['title']} id:{job['job_id']}"
								text = f"Profile Match: {match_percent}%"
								if job['formatted_work_type']:
									text = text + f", Work Type: {job['formatted_work_type']}"
								if job['min_salary']:
									text = text + f", Minimum Salary: {job['min_salary']}"
								if job['max_salary']:
									text = text + f", Maximum Salary: {job['max_salary']}"
								if job['pay_period']:
									text = text + f", Pay Period: {job['pay_period']}"

								options.append(
									{
										"type": "accordion",
										"title": job["title"],
										"subtitle": job["location"],
										"text": text
									})
								options.append({
										"type": "chips",
										"options": [
											{
											"text": option_text
											}
										]
									})
							json_response = {
								"page_info": {
											"form_info": {
												"parameter_info": [
													{
														"displayName": "results_displayed",
														"required": False,
														"state": "VALID",
														"value": True,
													},
												],
											},
										},
								'fulfillment_response': {
									'messages': [
										{"text": {"text": ["Here are a few matches, please click on 'Export' to save the detailed descriptions."]}},
										{
											'payload': {
												'richContent': [options]
											}
										}
									]
								}
							}

				return jsonify(json_response)
			if tag == "job_export":
				job_id = request_json["text"].split("id:")[1]
				job_name = request_json["text"][8:]
				job_details = get_job(job_id)
				matches_folder_id = session_parameters["matches_folder_id"]
				matches_folder_link = session_parameters["matches_folder_link"]
				upload_success = save_job(job_details,job_name,matches_folder_id)
				if upload_success:
					html =  f'''
					<p>The job details for {job_name} have been saved to Google Drive.</p>
					<p><a href="{matches_folder_link}" target="_blank">Access Link</a></p>
					<p>I can also draft a cover letter for any of the exported jobs.</p>
					'''
					json_response = {
							'fulfillment_response': {
								'messages': [
									{
										'payload': {
											'richContent': [
												[
													{
														"type": "html",
														"html": html
													},
													{
														"type": "chips",
														"options": [
															{
															"text": "Help me create a cover letter."
															}
														]
													}
												]
											]
										}
									}
								]
							}
						}
					return jsonify(json_response)
			if tag == "cl_select_resume":
				resume_folder_id = session_parameters["resume_folder_id"]
				files = get_folder_contents(resume_folder_id)
				if len(files) == 0:
					text = "I was unable to find any Resumes."
					json_response = {
							'fulfillment_response': {
								'messages': [
									{"text": {"text": [text]}}
								]
							}
						}
				else:
					options = []
					text = "Please click on the file name to select the Resume."
					for file in files:
						option_text = "Filename: "+file["name"]
						options.append({
							"type": "chips",
							"options": [
								{
								"text": option_text
								}
							]
						})
					json_response = {
							'fulfillment_response': {
								'messages': [
									{"text": {"text": [text]}},
									{
										'payload': {
											'richContent': [options]
										}
									}
								]
							}
						}
				return jsonify(json_response)
			if tag == "cl_select_job":
				matches_folder_id = session_parameters["matches_folder_id"]
				files = get_folder_contents(matches_folder_id)
				if len(files) == 0:
					text = "I was unable to find any matching jobs."
					json_response = {
							'fulfillment_response': {
								'messages': [
									{"text": {"text": [text]}}
								]
							}
						}
				else:
					options = []
					text = "Please click on the file name to select the job."
					for file in files:
						option_text = "Filename: "+file["name"]
						options.append({
							"type": "chips",
							"options": [
								{
								"text": option_text
								}
							]
						})
					json_response = {
							'fulfillment_response': {
								'messages': [
									{"text": {"text": [text]}},
									{
										'payload': {
											'richContent': [options]
										}
									}
								]
							}
						}
				return jsonify(json_response)
			if tag == "create_coverletter":
				matches_folder_id = session_parameters["matches_folder_id"]
				resume_folder_id = session_parameters["resume_folder_id"]
				cl_folder_id = session_parameters["cl_folder_id"]
				cl_folder_link = session_parameters["cl_folder_link"]

				for parameter in page_parameters:
					if parameter["displayName"] == "resume":
						resume_name = parameter["value"]
						resume_name = resume_name[10:]
					if parameter["displayName"] == "job":
						job_name = parameter["value"]
						job_name = job_name[10:]

				resume_path = download_file(resume_folder_id, resume_name)
				if resume_path and resume_path.endswith(".docx"):
					resume_text = get_txt_docx(resume_path)
				if resume_path and resume_path.endswith(".pdf"):
					resume_text = get_txt_pdf(resume_path)

				job_path = download_file(matches_folder_id, job_name)
				if job_path and job_path.endswith(".docx"):
					job_text = get_txt_docx(job_path)
				if job_path and job_path.endswith(".pdf"):
					job_text = get_txt_pdf(job_path)

				cl_text = generate_cover_letter(resume_text,job_text)
				cl_file_name = job_name.split(".docx")[0]
				upload_success = save_cl(cl_text,cl_file_name,cl_folder_id)
				if upload_success:
					html =  f'''
					<p>The cover letter for {job_name} has been saved to Google Drive.</p>
					<p><a href="{cl_folder_link}" target="_blank">Access Link</a></p>
					'''
					json_response = {
							'fulfillment_response': {
								'messages': [
									{
										'payload': {
											'richContent': [
												[
													{
														"type": "html",
														"html": html
													}
												]
											]
										}
									}
								]
							}
						}
					return jsonify(json_response)
			if tag == "delete_folders":
				session_folder_id = session_parameters["session_folder_id"]
				folders_deleted = delete_folders(session_folder_id)
				if folders_deleted:
					json_response = {
							'fulfillment_response': {
								'messages': [
									{"text": {"text": ["All the data has been deleted! Hope to hear from you again."]}},
								]
							}
						}
					return jsonify(json_response)

	return 'OK'